#!/usr/bin/env python3
"""
Builds skeleton .docx deliverables for a TOGAF ADM phase, pre-filled with
questionnaire answers where relevant. Requires python-docx (pip install python-docx).
"""

from datetime import date

from docx import Document

# Which answer keys are relevant "Scope" bullets per document key.
# Falls back to a generic set if a doc key isn't listed here.
DOC_SCOPE_FIELDS = {
    "01_Architecture_Vision": ["engagement_type", "industry", "company_size", "geo_scope", "drivers"],
    "A_Statement_of_Architecture_Work": ["engagement_type", "drivers"],
    "A_Capability_Assessment": ["business_processes", "current_erp"],
    "A_Communications_Plan": ["company_size", "geo_scope"],
    "A_Architecture_Definition_Document": ["drivers", "principles"],
    "02_Business_Architecture": ["business_processes"],
    "03_Data_Architecture": ["applications", "business_processes"],
    "04_Application_Architecture": ["applications"],
    "CD_Architecture_Deliverables_Phases_C_and_D": ["applications", "platform_primary"],
    "05_Technology_Architecture": ["platform_primary", "btp_services", "hyperscaler"],
    "06_Architecture_Requirements_Specification": ["pain_points", "drivers"],
    "E_Architecture_Building_Blocks": ["applications", "platform_primary"],
    "E_Solution_Building_Blocks": ["applications", "btp_services"],
    "07_Implementation_and_Migration_Plan": ["current_erp", "pain_points"],
    "F_Architecture_Contract": ["applications"],
    "08_Architecture_Roadmap": ["drivers", "pain_points"],
    "G_Implementation_Governance_Model": ["principles"],
    "G_Compliance_Assessment": ["pain_points"],
    "H_Requirements_Impact_Assessment": ["pain_points"],
    "H_Change_Request": ["pain_points"],
    "Prelim_Request_for_Architecture_Work": ["engagement_type", "drivers"],
    "Prelim_Organizational_Model_for_EA": ["company_size", "geo_scope"],
    "Prelim_Tailored_Architecture_Framework": ["principles"],
    "Prelim_Architecture_Repository": ["principles"],
    "Prelim_Business_Principles_Goals_and_Drivers": ["principles", "drivers"],
    "Prelim_Architecture_Principles": ["principles"],
    "E_Opportunities_and_Solutions_Assessment": ["business_processes", "applications", "pain_points"],
    "E_Architecture_Definition_Document": ["drivers", "principles", "applications", "platform_primary"],
}

FIELD_LABELS = {
    "engagement_type": "Engagement type",
    "industry": "Industry sector",
    "company_size": "Company size",
    "geo_scope": "Geographic scope",
    "current_erp": "Current ERP / core system landscape",
    "pain_points": "Current issues / pain points",
    "drivers": "Strategic drivers",
    "principles": "Architecture principles",
    "business_processes": "Business processes in scope",
    "applications": "Proposed applications",
    "platform_primary": "Primary platform / hosting model",
    "btp_services": "Supporting SAP BTP services",
    "hyperscaler": "Hyperscaler preference",
}

PURPOSE_TEXT = {
    "01_Architecture_Vision": "Describes the high-level vision of the target enterprise architecture, its business value, and the scope of the engagement.",
    "02_Business_Architecture": "Describes the target business architecture: business processes, actors, roles and organizational structure.",
    "03_Data_Architecture": "Describes the logical and physical data assets and data management resources required to support the business.",
    "04_Application_Architecture": "Describes the application components and their interactions required to support the business processes.",
    "05_Technology_Architecture": "Describes the technology platform, infrastructure and hosting model required to support the applications.",
    "06_Architecture_Requirements_Specification": "Consolidates the architecture requirements gathered across Phases B, C and D.",
    "07_Implementation_and_Migration_Plan": "Describes the approach to move from the baseline to the target architecture, including work packages and sequencing.",
    "08_Architecture_Roadmap": "Lists the individual work packages / projects in a time-sequenced roadmap toward the target architecture.",
    "Prelim_Architecture_Principles": "States the formal architecture principles that will govern decision-making throughout the engagement, separate from the business principles/goals/drivers captured alongside them.",
    "E_Opportunities_and_Solutions_Assessment": "Assesses candidate opportunities and solution building blocks against the target architecture, grouping them into work packages and informing the transition roadmap.",
    "E_Architecture_Definition_Document": "Consolidates and formally baselines the Architecture Definition Document across the Business, Data, Application and Technology architectures defined in Phases B-D, superseding the Phase A draft.",
}
DEFAULT_PURPOSE = "Supporting TOGAF ADM artifact for this phase. Elaborate with stakeholders during the corresponding ADM workshop."


def _bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def build_doc(path, doc_key: str, doc_title: str, phase_title: str, answers: dict, short_code: str):
    doc = Document()

    doc.add_heading(doc_title, level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Repository: {answers.get('repository_name', '')}\n").italic = True
    meta.add_run(f"Company: {answers.get('company_name', '')}\n").italic = True
    meta.add_run(f"TOGAF ADM phase: {phase_title}\n").italic = True
    meta.add_run(f"Author: {answers.get('author', '')}\n").italic = True
    meta.add_run(f"Generated: {date.today().isoformat()} (base template)").italic = True

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(PURPOSE_TEXT.get(doc_key, DEFAULT_PURPOSE))

    doc.add_heading("Scope", level=1)
    scope_fields = DOC_SCOPE_FIELDS.get(doc_key, ["drivers", "pain_points"])
    any_scope = False
    for field in scope_fields:
        value = answers.get(field)
        if not value:
            continue
        any_scope = True
        doc.add_paragraph(FIELD_LABELS.get(field, field), style="Heading 3")
        if isinstance(value, list):
            _bullets(doc, value)
        else:
            doc.add_paragraph(str(value))
    if not any_scope:
        doc.add_paragraph("To be defined with stakeholders.")

    doc.add_heading("Content", level=1)
    doc.add_paragraph(
        "Draft placeholder generated from the base template. Replace this section with the "
        "actual deliverable content produced during the corresponding TOGAF ADM phase workshop."
    )

    doc.add_heading("Status", level=1)
    status = doc.add_paragraph()
    status.add_run("DRAFT").bold = True
    status.add_run(f" - generated from base template on {date.today().isoformat()}.")

    doc.save(str(path))
